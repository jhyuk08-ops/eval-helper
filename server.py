import os
import json
import re
import unicodedata
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai

app = Flask(__name__)

# =========================================================
# 상대 경로 설정 (BASE_DIR 기준)
# =========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 성취기준(평가기준) JSON 저장 경로
EVAL_STD_MIDDLE_DIR = os.path.join(BASE_DIR, "data", "evaluation standard", "middle")
EVAL_STD_HIGH_DIR = os.path.join(BASE_DIR, "data", "evaluation standard", "high")

# 성취수준 JSON 저장 경로
ACHIEVE_LEVEL_MIDDLE_DIR = os.path.join(BASE_DIR, "data", "achievement level", "middle")
ACHIEVE_LEVEL_HIGH_DIR = os.path.join(BASE_DIR, "data", "achievement level", "high")


def normalize_str(s):
    """한글 자음/모음 분리 현상(NFD) 방지 및 공백 정규화"""
    if s is None:
        return ""
    return unicodedata.normalize('NFC', str(s)).strip()


def extract_code(text):
    """
    텍스트에서 [코드] 추출 (예: '[9사01-01] 사회적 현상...' -> '9사01-01')
    대괄호가 없을 경우 문자열 전체 반환
    """
    if not text:
        return ""
    text_str = normalize_str(text)
    match = re.search(r'\[(.*?)\]', text_str)
    if match:
        return match.group(1).strip()
    return text_str


def set_cell_background(cell, fill_hex):
    """표 셀 배경색 지정 함수"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_vertical_align(cell, align="center"):
    """표 셀 세로 정렬 지정 함수"""
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def group_achievement_levels(levels):
    """
    동일한 성취수준 내용을 가진 등급들을 'A, B' 형태로 그룹화합니다.
    """
    if not levels:
        return []

    grouped = [] # [{'grades': ['A', 'B'], 'desc': '...'}, ...]
    
    for lvl in levels:
        grade = lvl.get('grade', '').strip()
        desc = lvl.get('desc', '').strip()

        if not desc:
            continue

        # 기존 항목 중 내용이 완전히 동일한 항목이 있는지 확인
        found = False
        for g_item in grouped:
            if g_item['desc'] == desc:
                if grade and grade not in g_item['grades']:
                    g_item['grades'].append(grade)
                found = True
                break

        if not found:
            grouped.append({
                'grades': [grade] if grade else [],
                'desc': desc
            })

    # 'A, B' 형태로 문자열 변환 후 리스트 구성
    result = []
    for g_item in grouped:
        grade_str = ", ".join(g_item['grades']) if g_item['grades'] else "-"
        result.append({
            'grade': grade_str,
            'desc': g_item['desc']
        })

    return result


def read_achievement_levels(school_type, main_subject, detail_subject=""):
    """
    data/achievement level/{middle|high}/{교과명}.json 경로에서 성취수준 읽기
    """
    target_dir = ACHIEVE_LEVEL_HIGH_DIR if school_type == 'high' else ACHIEVE_LEVEL_MIDDLE_DIR
    
    file_path = os.path.join(target_dir, f"{main_subject}.json")
    if detail_subject and not os.path.exists(file_path):
        alt_path = os.path.join(target_dir, f"{detail_subject}.json")
        if os.path.exists(alt_path):
            file_path = alt_path

    standards_map = {}

    if not os.path.exists(file_path):
        print(f"[성취수준 경고] 파일이 존재하지 않습니다: {file_path}")
        return standards_map

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    raw_code = item.get('성취기준(코드)', item.get('성취기준 코드', item.get('code', '')))
                    code_key = extract_code(raw_code)

                    grade = str(item.get('성취수준(등급)', item.get('성취수준 등급', item.get('grade', '')))).strip()
                    desc = str(item.get('성취수준(내용)', item.get('성취수준 내용', item.get('desc', '')))).strip()

                    if not code_key:
                        continue

                    if code_key not in standards_map:
                        standards_map[code_key] = []

                    if grade or desc:
                        standards_map[code_key].append({
                            'grade': grade,
                            'desc': desc
                        })

        print(f"[성취수준 성공] {os.path.basename(file_path)} 에서 {len(standards_map)}개 성취기준 읽기 완료")
    except Exception as e:
        print(f"[성취수준 읽기 오류]: {e}")

    return standards_map


def read_standards_from_json(file_path):
    """evaluation standard JSON 파일 읽기"""
    if not os.path.exists(file_path):
        return {}

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        standards_map = {}
        if isinstance(data, list):
            for item in data:
                if isinstance(item, str):
                    std_str = normalize_str(item)
                    code = extract_code(std_str)
                    standards_map[code] = {'full_text': std_str}
                elif isinstance(item, dict):
                    code = str(item.get('성취기준(코드)', item.get('성취기준 코드', item.get('code', '')))).strip()
                    name = str(item.get('성취기준명', item.get('성취기준 내용', item.get('name', '')))).strip()

                    full_text = str(item.get('성취기준(코드, 내용)', item.get('full_text', ''))).strip()
                    if not full_text:
                        if code and name: full_text = f"[{code}] {name}"
                        elif name: full_text = name
                        elif code: full_text = f"[{code}]"

                    if not full_text: continue

                    code_key = extract_code(full_text) or code
                    standards_map[code_key] = {
                        'code': code_key,
                        'full_text': full_text
                    }
        return standards_map
    except Exception as e:
        print(f"[평가기준 JSON 읽기 오류]: {e}")
        return {}


def get_json_file_path(school_type, main_subject, detail_subject):
    main_subject = normalize_str(main_subject)
    detail_subject = normalize_str(detail_subject)
    if school_type == 'middle':
        return os.path.join(EVAL_STD_MIDDLE_DIR, f"{main_subject}.json")
    else:
        return os.path.join(EVAL_STD_HIGH_DIR, main_subject, f"{detail_subject}.json")


# =========================================================
# 라우트 핸들러
# =========================================================

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/get_subjects', methods=['POST'])
def get_subjects():
    try:
        data = request.get_json() or {}
        school_type = data.get('schoolType', 'high')

        if school_type == 'middle':
            subjects = []
            if os.path.exists(EVAL_STD_MIDDLE_DIR):
                for fname in os.listdir(EVAL_STD_MIDDLE_DIR):
                    if fname.lower().endswith('.json'):
                        subjects.append(normalize_str(os.path.splitext(fname)[0]))
                subjects.sort()
            return jsonify({'success': True, 'subjects': subjects})
        else:
            hierarchy = {}
            if os.path.exists(EVAL_STD_HIGH_DIR):
                for folder_name in os.listdir(EVAL_STD_HIGH_DIR):
                    folder_path = os.path.join(EVAL_STD_HIGH_DIR, folder_name)
                    if os.path.isdir(folder_path):
                        norm_folder = normalize_str(folder_name)
                        json_files = []
                        for fname in os.listdir(folder_path):
                            if fname.lower().endswith('.json'):
                                json_files.append(normalize_str(os.path.splitext(fname)[0]))
                        json_files.sort()
                        hierarchy[norm_folder] = json_files
            return jsonify({'success': True, 'subjects': sorted(list(hierarchy.keys())), 'hierarchy': hierarchy})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/get_standards', methods=['POST'])
def get_standards():
    try:
        data = request.get_json() or {}
        school_type = data.get('schoolType', 'high')
        main_subject = data.get('교과', '')
        detail_subject = data.get('세부교과', '')

        file_path = get_json_file_path(school_type, main_subject, detail_subject)
        json_map = read_standards_from_json(file_path)

        standards_list = []
        for k, v in json_map.items():
            if isinstance(v, dict) and v.get('full_text'):
                standards_list.append(v['full_text'])
            else:
                standards_list.append(k)

        return jsonify({
            'success': True,
            'standards': standards_list
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/generate', methods=['POST'])
def generate():
    try:
        api_key = request.form.get('apiKey')
        school_type = request.form.get('schoolType', 'high')
        school_type_str = "고등학교" if school_type == 'high' else "중학교"
        grade = request.form.get('grade', '1')
        semester = request.form.get('semester', '1학기')
        main_subject = request.form.get('교과', '')
        detail_subject = request.form.get('세부교과', '')

        selected_standards = json.loads(request.form.get('selectedStandards', '[]'))
        exam_ratios = json.loads(request.form.get('examRatios', '[0, 0]'))
        exam_standards = json.loads(request.form.get('examStandards', '{}'))
        raw_perf_list = json.loads(request.form.get('perfList', '[]'))

        # 성취수준 JSON 읽기
        achieve_levels_map = read_achievement_levels(school_type, main_subject, detail_subject)

        achieve_ref_rubrics = []
        std_levels_dict = {}

        for std_text in selected_standards:
            code = extract_code(std_text)
            raw_levels = achieve_levels_map.get(code, [])
            
            # 내용이 같은 성취수준 등급 그룹화 (예: A, B / C, D)
            grouped_levels = group_achievement_levels(raw_levels)
            std_levels_dict[std_text] = grouped_levels

            achieve_ref_rubrics.append({
                "code_and_name": std_text,
                "code": code,
                "levels": grouped_levels
            })

        perf_prompt_structure = []
        for perf in raw_perf_list:
            d_name = perf.get('domain', '')
            sub_items = perf.get('subItems', [])
            processed_subs = []
            for sub in sub_items:
                s_name = sub.get('name', '')
                try:
                    s_max = int(sub.get('max', 0))
                    s_min = int(sub.get('min', 0))
                    s_step = int(sub.get('step', 1))
                    if s_step <= 0: s_step = 1
                    scores = list(range(s_max, s_min - 1, -s_step))
                    if not scores: scores = [s_max]
                except Exception:
                    scores = [sub.get('max', 0)]

                processed_subs.append({
                    "sub_name": s_name,
                    "scores": scores
                })
            perf_prompt_structure.append({
                "domain_name": d_name,
                "sub_items": processed_subs
            })

        # Gemini AI 생성
        client = genai.Client(api_key=api_key)
        prompt = f"""
당신은 대한민국 교육과정 및 학생평가 전문가입니다.
다음 데이터를 바탕으로 평가계획서의 주요 내용을 작성해주세요.

[기본 정보]
- 학교급 및 학년: {school_type_str} {grade}학년 ({semester})
- 교과/세부교과: {main_subject} / {detail_subject if school_type == 'high' else '없음'}
- 선택된 성취기준: {json.dumps(selected_standards, ensure_ascii=False)}
- 성취수준 참조 데이터: {json.dumps(achieve_ref_rubrics, ensure_ascii=False)}
- 수행평가 배점 구조: {json.dumps(perf_prompt_structure, ensure_ascii=False)}

[작성 요구사항 - 엄격 준수]
1. purpose: 교과 특성에 맞는 평가 목적 5개 항목("가.", "나.", "다.", "라.", "마.")
2. direction: 평가 기본 방향 및 방침 ("가. 기본 방향: ...", "나. 방침: ...")
3. rubrics: 수행평가 영역별 대표 루브릭 문장 (상, 중, 하)
4. sub_item_details: 수행평가 요소별로 계산된 각 점수(예: 3점, 2점, 1점 등)에 해당하는 세부 평가 내용(채점 기준)을 명확하고 구체적으로 작성해 주세요.

[응답 포맷] 순수 JSON만 반환 (Markdown 블록 제외)
{{
  "purpose": [
    "가. ...",
    "나. ...",
    "다. ...",
    "라. ...",
    "마. ..."
  ],
  "direction": [
    "가. 기본 방향: ...",
    "나. 방침: ..."
  ],
  "rubrics": [
    {{
      "domain_name": "(영역명)",
      "criteria_high": "(상)",
      "criteria_mid": "(중)",
      "criteria_low": "(하)"
    }}
  ],
  "sub_item_details": [
    {{
      "domain_name": "(영역명)",
      "sub_name": "(요소명)",
      "score_criteria": [
        {{"score": 3, "description": "세부 평가 기준 내용..."}},
        {{"score": 2, "description": "세부 평가 기준 내용..."}},
        {{"score": 1, "description": "세부 평가 기준 내용..."}}
      ]
    }}
  ]
}}
"""

        candidate_models = ['gemini-3.6-flash', 'gemini-2.5-flash']
        response = None
        last_err = None

        for m_name in candidate_models:
            try:
                response = client.models.generate_content(
                    model=m_name,
                    contents=prompt
                )
                if response and response.text:
                    break
            except Exception as e_mod:
                last_err = e_mod

        if not response:
            raise last_err

        try:
            raw_text = response.text.strip().replace('```json', '').replace('```', '')
            ai_data = json.loads(raw_text)
        except Exception as e:
            print(f"[AI 응답 파싱 에러]: {e}")
            ai_data = {
                "purpose": [
                    f"가. {main_subject} 교과 교육과정 성취기준에 부합하는 종합적 사고력 평가",
                    "나. 학습 과정과 결과를 종합적으로 평가하여 학생의 성장을 지원함",
                    "다. 교과 역량 및 성취 수준을 공정하게 측정함",
                    "라. 평가 결과를 수업 개선 및 맞춤형 지도 자료로 활용함",
                    "마. 자기주도적 학습 태도 함양 도모"
                ],
                "direction": [
                    "가. 기본 방향: 지필평가와 수행평가를 균형 있게 실시함.",
                    "나. 방침: 평가 계획을 공지하고 공정성과 신뢰성을 확보함."
                ],
                "rubrics": [],
                "sub_item_details": []
            }

        # Word 문서 생성
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        title_subject = detail_subject if (school_type == 'high' and detail_subject) else main_subject

        heading = doc.add_heading(f"2026학년도 {title_subject}과 평가계획서", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # 1. 기본 정보
        doc.add_heading("1. 기본 정보", level=1)
        doc.add_heading("가. 학교급 및 학년", level=2)
        p1 = doc.add_paragraph()
        p1.paragraph_format.left_indent = Inches(0.2)
        p1.add_run(f"• 학교급: {school_type_str}\n")
        p1.add_run(f"• 학년 및 학기: {grade}학년 {semester}\n")
        p1.add_run(f"• 교과: {main_subject}" + (f" / 세부교과: {detail_subject}" if school_type == 'high' else ""))

        # 나. 성취기준 및 성취수준 표
        doc.add_heading("나. 성취기준 및 성취수준", level=2)
        if selected_standards:
            std_table = doc.add_table(rows=1, cols=3)
            std_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            std_table.style = 'Table Grid'

            hdr_cells = std_table.rows[0].cells
            hdr_titles = ["성취기준 (코드 및 내용)", "성취수준 (등급)", "성취수준 (내용)"]
            widths = [Inches(2.5), Inches(1.0), Inches(3.3)]

            for i, title in enumerate(hdr_titles):
                hdr_cells[i].text = title
                hdr_cells[i].width = widths[i]
                set_cell_background(hdr_cells[i], 'F2F2F2')
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for std in selected_standards:
                levels = std_levels_dict.get(std, [])

                if levels:
                    start_row_idx = len(std_table.rows)
                    for lvl in levels:
                        row_cells = std_table.add_row().cells
                        row_cells[0].text = std
                        row_cells[1].text = lvl.get('grade', '-') # 'A, B' 혹은 'A'
                        row_cells[2].text = lvl.get('desc', '')

                        for idx, cell in enumerate(row_cells):
                            cell.width = widths[idx]

                        # 등급 셀 중앙 정렬
                        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    end_row_idx = len(std_table.rows) - 1

                    # 성취기준 셀 1번만 표기되도록 세로 병합 및 세로 중앙 정렬
                    if end_row_idx > start_row_idx:
                        top_cell = std_table.cell(start_row_idx, 0)
                        bottom_cell = std_table.cell(end_row_idx, 0)
                        top_cell.merge(bottom_cell)
                        top_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_vertical_align(top_cell, "center")
                    elif start_row_idx <= end_row_idx:
                        c = std_table.cell(start_row_idx, 0)
                        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_vertical_align(c, "center")
                else:
                    row_cells = std_table.add_row().cells
                    row_cells[0].text = std
                    row_cells[1].text = "-"
                    row_cells[2].text = "성취수준 내용이 없습니다."
                    for idx, cell in enumerate(row_cells):
                        cell.width = widths[idx]
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p_none = doc.add_paragraph("• 선택된 성취기준이 없습니다.")
            p_none.paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

        # 다. 정기시험
        doc.add_heading("다. 정기시험 연동 성취기준 및 반영비율", level=2)
        p_exam = doc.add_paragraph()
        p_exam.paragraph_format.left_indent = Inches(0.2)
        p_exam.add_run(f"• 1차 지필평가 (반영비율: {exam_ratios[0]}%)\n")
        p_exam.add_run(f"  - 연동 성취기준: {', '.join(exam_standards.get('exam1', [])) or '선택 없음'}\n\n")
        p_exam.add_run(f"• 2차 지필평가 (반영비율: {exam_ratios[1]}%)\n")
        p_exam.add_run(f"  - 연동 성취기준: {', '.join(exam_standards.get('exam2', [])) or '선택 없음'}")

        doc.add_paragraph()

        # 2. 평가의 목적
        doc.add_heading("2. 평가의 목적", level=1)
        for item in ai_data.get("purpose", []):
            doc.add_paragraph(item).paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

        # 3. 평가의 기본 방향과 방침
        doc.add_heading("3. 평가의 기본 방향과 방침", level=1)
        for item in ai_data.get("direction", []):
            doc.add_paragraph(item).paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

        # 4. 수행평가 세부기준
        doc.add_heading("4. 수행평가 세부기준", level=1)
        ai_sub_details = ai_data.get('sub_item_details', [])

        if raw_perf_list:
            for idx, perf in enumerate(raw_perf_list, 1):
                domain_name = perf.get('domain', f'영역 {idx}')
                doc.add_heading(f"가. [영역 {idx}] {domain_name} (반영비율: {perf.get('ratio', 0)}%, 기본점수: {perf.get('baseScore', 0)}점)", level=2)

                p_perf_stds = doc.add_paragraph()
                p_perf_stds.paragraph_format.left_indent = Inches(0.2)
                p_perf_stds.add_run(f"• 연동 성취기준: {', '.join(perf.get('standards', [])) or '선택 없음'}\n")

                sub_items = perf.get('subItems', [])
                if sub_items:
                    doc.add_paragraph("• 평가요소별 점수 산정 및 세부 평가기준").paragraph_format.left_indent = Inches(0.2)
                    table = doc.add_table(rows=1, cols=3)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'

                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text = "평가 요소명", "배점", "세부 평가 내용 (채점 기준)"
                    table.columns[0].width = Inches(1.5)
                    table.columns[1].width = Inches(1.0)
                    table.columns[2].width = Inches(4.3)

                    for h_cell in hdr:
                        set_cell_background(h_cell, 'F2F2F2')
                        h_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for sub in sub_items:
                        s_name = sub.get('name', '')
                        ai_criteria_item = next(
                            (item for item in ai_sub_details if item.get('domain_name') == domain_name and item.get('sub_name') == s_name),
                            None
                        )

                        score_c_list = ai_criteria_item.get('score_criteria', []) if ai_criteria_item else []

                        if score_c_list:
                            for sc in score_c_list:
                                row = table.add_row().cells
                                row[0].text = s_name
                                row[1].text = f"{sc.get('score')}점"
                                row[2].text = str(sc.get('description', ''))
                                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            try:
                                s_max = int(sub.get('max', 0))
                                s_min = int(sub.get('min', 0))
                                s_step = int(sub.get('step', 1))
                                if s_step <= 0: s_step = 1
                                scores = list(range(s_max, s_min - 1, -s_step))
                            except Exception:
                                scores = [sub.get('max', 0)]

                            for sc_val in scores:
                                row = table.add_row().cells
                                row[0].text = s_name
                                row[1].text = f"{sc_val}점"
                                row[2].text = f"{s_name} 영역 성취 수준에 따른 {sc_val}점 세부 채점 기준 적용"
                                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc.add_paragraph()

                rubric = next((r for r in ai_data.get('rubrics', []) if r.get('domain_name') == domain_name), None)
                doc.add_paragraph("• 영역별 총괄 성취수준 (루브릭)").paragraph_format.left_indent = Inches(0.2)
                r_table = doc.add_table(rows=4, cols=2)
                r_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                r_table.style = 'Table Grid'
                r_table.columns[0].width, r_table.columns[1].width = Inches(1.2), Inches(5.6)

                r_table.cell(0, 0).text, r_table.cell(0, 1).text = "성취수준", "평가 기준 내용"
                set_cell_background(r_table.cell(0, 0), 'F2F2F2')
                set_cell_background(r_table.cell(0, 1), 'F2F2F2')

                r_table.cell(1, 0).text, r_table.cell(1, 1).text = "상", rubric.get('criteria_high', '성취기준을 우수하게 달성함') if rubric else '상'
                r_table.cell(2, 0).text, r_table.cell(2, 1).text = "중", rubric.get('criteria_mid', '성취기준을 보통으로 달성함') if rubric else '중'
                r_table.cell(3, 0).text, r_table.cell(3, 1).text = "하", rubric.get('criteria_low', '성취기준 달성이 노력 요함') if rubric else '하'

                for r_idx in range(1, 4):
                    r_table.cell(r_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

        now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        if school_type == 'high' and detail_subject and detail_subject != main_subject:
            subj_filename_part = f"{main_subject}({detail_subject})"
        else:
            subj_filename_part = main_subject

        filename = f"{school_type_str}_{grade}학년_{semester}_{subj_filename_part}_{now_str}.docx"

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        print(f"[ERROR] 문서 생성 에러: {e}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print(f"서버 실행 중... BASE_DIR: {BASE_DIR}")
    app.run(debug=True, host='0.0.0.0', port=5000)