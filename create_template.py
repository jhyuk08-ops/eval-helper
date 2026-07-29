import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def create_default_template():
    doc = docx.Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 1. 문서 제목
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("2022 개정 교육과정 {{ subject_name }}과 평가계획서")
    run_title.font.name = "맑은 고딕"
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    doc.add_paragraph()

    # 1. 기본 정보
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. 기본 정보")
    r_h1.font.name = "맑은 고딕"
    r_h1.font.size = Pt(12)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    t1 = doc.add_table(rows=2, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    t1_data = [
        [("학교급 및 학년", True), ("{{ school_info }}", False)],
        [("지필평가 비율", True), ("{{ paper_info }}", False)]
    ]

    for r_idx, row in enumerate(t1_data):
        for c_idx, (text, is_header) in enumerate(row):
            cell = t1.cell(r_idx, c_idx)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.runs[0].font.name = "맑은 고딕"
            p.runs[0].font.size = Pt(10)
            
            if is_header:
                set_cell_background(cell, "F1F5F9")
                p.runs[0].font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # 2. 평가의 목적
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. 평가의 목적")
    r_h2.font.name = "맑은 고딕"
    r_h2.font.size = Pt(12)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p_purposes = doc.add_paragraph()
    r_purposes = p_purposes.add_run("{{ eval_purpose }}")
    r_purposes.font.name = "맑은 고딕"
    r_purposes.font.size = Pt(10)

    doc.add_paragraph()

    # 3. 평가의 기본 방향과 방침
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. 평가의 기본 방향과 방침")
    r_h3.font.name = "맑은 고딕"
    r_h3.font.size = Pt(12)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # 가. 기본 방향
    p_dir_head = doc.add_paragraph()
    r_dir_head = p_dir_head.add_run("  가. 기본 방향")
    r_dir_head.font.name = "맑은 고딕"
    r_dir_head.font.size = Pt(10.5)
    r_dir_head.font.bold = True

    p_dir_body = doc.add_paragraph()
    r_dir_body = p_dir_body.add_run("{{ eval_direction }}")
    r_dir_body.font.name = "맑은 고딕"
    r_dir_body.font.size = Pt(10)

    # 나. 방침
    p_pol_head = doc.add_paragraph()
    r_pol_head = p_pol_head.add_run("  나. 방침")
    r_pol_head.font.name = "맑은 고딕"
    r_pol_head.font.size = Pt(10.5)
    r_pol_head.font.bold = True

    policies = [
        "    1) 성적반영 비율은 지필평가 {{ paper_ratio_sum }}%, 수행평가 {{ perf_ratio_sum }}%로 한다.",
        "    2) 지필평가는 {{ paper_count }}회 실시한다.",
        "    3) 논술형 평가는 수업 중에 실시하며, 단편적인 지식이나 하나의 정답만을 요구하기보다는 학생들이 생각한 바를 논리적으로 조직하여 표현하는 능력과 학습 내용에 대한 종합적인 이해와 고차원적인 학습 결과를 측정하도록 한다.",
        "    4) 논술형 평가 문제 출제 시 구체적인 채점기준안을 마련하여 공동 채점한다. 단, 논술형 평가에서 인정 답안에 제시된 답안의 범위에서 벗어나는 경우 교과협의를 거쳐 객관적이고 신뢰성 있는 평가가 되도록 한다.",
        "    5) 동 학년 교사들 간에 평가 협의를 통하여 다양한 평가 방법과 기법 및 도구를 활용하여 객관적이고도 신뢰성 있는 평가를 실시한다.",
        "    6) 평가 결과는 학생들의 전인적인 발달을 도모하고 교수․학습을 개선하는데 필요한 자료로 활용할 수 있도록 한다."
    ]

    for pol in policies:
        p_pol = doc.add_paragraph()
        r_pol = p_pol.add_run(pol)
        r_pol.font.name = "맑은 고딕"
        r_pol.font.size = Pt(10)

    doc.add_paragraph()

    # 4. 수행평가 세부 영역 및 성취수준
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. 수행평가 세부 영역 및 성취수준")
    r_h4.font.name = "맑은 고딕"
    r_h4.font.size = Pt(12)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    t2 = doc.add_table(rows=2, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["수행평가 영역명 (비율)", "성취기준", "성취수준 (상 / 중 / 하)"]
    for c_idx, text in enumerate(headers):
        cell = t2.cell(0, c_idx)
        cell.text = text
        set_cell_background(cell, "E2E8F0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = "맑은 고딕"
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.bold = True

    cell_0 = t2.cell(1, 0)
    p0 = cell_0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.text = "{% for item in perf_evals %}{{ item.name }}\n({{ item.ratio }}%)"
    p0.runs[0].font.name = "맑은 고딕"
    p0.runs[0].font.size = Pt(9.5)
    p0.runs[0].font.bold = True

    cell_1 = t2.cell(1, 1)
    p1 = cell_1.paragraphs[0]
    p1.text = "{{ item.standard }}"
    p1.runs[0].font.name = "맑은 고딕"
    p1.runs[0].font.size = Pt(9)

    cell_2 = t2.cell(1, 2)
    p2 = cell_2.paragraphs[0]
    
    r_high = p2.add_run("[상] {{ item.eval_high }}\n\n")
    r_high.font.name = "맑은 고딕"
    r_high.font.size = Pt(8.5)
    r_high.font.bold = True
    r_high.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    r_mid = p2.add_run("[중] {{ item.eval_mid }}\n\n")
    r_mid.font.name = "맑은 고딕"
    r_mid.font.size = Pt(8.5)
    r_mid.font.bold = True
    r_mid.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    r_low = p2.add_run("[하] {{ item.eval_low }}")
    r_low.font.name = "맑은 고딕"
    r_low.font.size = Pt(8.5)
    r_low.font.bold = True
    r_low.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    r_end = p2.add_run("{% endfor %}")
    r_end.font.name = "맑은 고딕"
    r_end.font.size = Pt(1)

    doc.save("default_template.docx")
    print("✅ 요구사항이 적용된 새로운 'default_template.docx' 생성 완료!")

if __name__ == "__main__":
    create_default_template()